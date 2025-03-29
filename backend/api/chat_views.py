import logging
from django.utils import timezone
from django.core.cache import cache
from django.http import JsonResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from .models import Video
from .chat_models import ChatSession, ChatMessage
from django.views.generic import ListView, DetailView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, FileResponse
from django.contrib import messages
from io import BytesIO
import docx
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Set up logging
logger = logging.getLogger(__name__)

# Cache settings
SESSION_CACHE_TIMEOUT = 60 * 60 * 24 * 7  # Session cache retention for 7 days


def get_or_create_chat_session(request):
    """
    Get or create an active chat session
    Ensure only one active session is returned, and create a new one if none exists
    Note: New sessions will not automatically inherit any video associations
    Ensure each user has only one active session
    """
    # Check if session ID exists in the request
    session_id = request.session.get('current_chat_session_id')
    
    # Check if the current session is valid
    active_session = None
    if session_id:
        try:
            # Attempt to get an existing session, strictly checking if it's active
            active_session = ChatSession.objects.get(
                id=session_id, 
                user=request.user, 
                is_active=True,   # Ensure the session is active
                ended_at__isnull=True  # Ensure the session has no end time
            )
            logger.info(f"Found active session ID: {active_session.id} for user {request.user.username}")
        except ChatSession.DoesNotExist:
            # Session does not exist or has ended, clear the current session ID
            request.session.pop('current_chat_session_id', None)
            logger.info(f"Session {session_id} not found or inactive for user {request.user.username}, will check other active sessions")
    
    # If no active session is found, check for other active sessions
    if not active_session:
        # Find all active sessions for the user, sorted by creation time in reverse order (newest first)
        other_active_sessions = ChatSession.objects.filter(
            user=request.user,
            is_active=True,
            ended_at__isnull=True
        ).order_by('-created_at')
        
        if other_active_sessions.exists():
            # Use the most recent active session
            active_session = other_active_sessions.first()
            # Update the Django session with the session ID
            request.session['current_chat_session_id'] = active_session.id
            logger.info(f"Found another active session ID: {active_session.id} for user {request.user.username}")
    
    # If no active session is found, create a new one
    if not active_session:
        # Create a new session, but do not automatically inherit any video associations
        # Use raw SQL to create the session, avoiding potential signals or events
        from django.db import connection
        cursor = connection.cursor()
        
        # First, create the session record
        cursor.execute(
            """
            INSERT INTO api_chatsession 
            (user_id, created_at, ended_at, is_active, title, summary) 
            VALUES (%s, %s, %s, %s, %s, %s) 
            RETURNING id
            """, 
            [
                request.user.id, 
                timezone.now(), 
                None, 
                True, 
                f"Conversation {timezone.now().strftime('%Y-%m-%d %H:%M')}", 
                ""
            ]
        )
        
        # Get the ID of the newly inserted record
        new_session_id = cursor.fetchone()[0]
        
        # Get the newly created session object using the database ID
        active_session = ChatSession.objects.get(id=new_session_id)
        logger.info(f"Created brand new session ID: {active_session.id} for user {request.user.username} (no video associations)")
        
        # Save the session ID to the Django session
        request.session['current_chat_session_id'] = active_session.id
        
        logger.info(f"Created new chat session ID: {active_session.id} for user {request.user.username}")
    
    return active_session


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def end_chat_session(request):
    """End the current chat session and generate a summary, ensuring the next session is new (does not inherit video associations)"""
    session_id = request.session.get('current_chat_session_id')
    
    if not session_id:
        return JsonResponse({'status': 'error', 'message': 'No active session'})
    
    try:
        session = ChatSession.objects.get(id=session_id, user=request.user)
        
        # Generate the session summary
        summary = generate_ai_summary(session)
        
        # Update the session record
        session.is_active = False
        session.ended_at = timezone.now()
        session.summary = summary
        session.save()
        
        # Clear the current session ID - ensure the next session is new
        request.session.pop('current_chat_session_id', None)
        
        # Clear the cached session data, ensuring no video associations are inherited
        cache_key = f"chat_session:{request.user.id}"
        cache.delete(cache_key)
        
        # Prepare the response data
        result = {
            'status': 'success', 
            'message': 'Session ended and summary generated',
            'summary': summary,
            'session_id': session.id,
            'duration': session.get_duration(),
            'videos_count': session.get_videos_count(),
            'messages_count': session.get_messages_count(),
        }
        
        logger.info(f"User {request.user.username} ended chat session ID: {session.id}, duration: {session.get_duration()} minutes")
        return JsonResponse(result)
        
    except ChatSession.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Session does not exist'})


def generate_ai_summary(session):
    """Generate a session summary using AI"""
    try:
        # Get all messages in the session
        messages = session.messages.all().order_by('timestamp')
        
        if not messages:
            return "No conversation content in this session."
        
        # Get the titles of associated videos
        video_titles = [video.title for video in session.session_videos.all() if video.title]
        video_info = ""
        if video_titles:
            video_info = f"Videos discussed in this session: {', '.join(video_titles)}\n\n"
        
        # Prepare the prompt
        prompt = "Please generate a concise summary of the conversation based on the following points:\n"
        prompt += "1. Main topics and content discussed\n"
        prompt += "2. Important language learning points (vocabulary, expressions, grammar, etc.)\n"
        prompt += "3. Cultural background knowledge\n"
        prompt += "The summary should be brief and highlight the most valuable information.\n\n"
        
        if video_info:
            prompt += video_info
        
        # Add the conversation content
        for msg in messages:
            role = "User" if msg.role == "user" else "AI Assistant"
            prompt += f"{role}: {msg.content}\n\n"
        
        # Call the AI model, importing it dynamically to avoid circular references
        try:
            # Dynamically import the contents of gemini_views
            from .gemini_views import client, GEMINI_MODEL
            
            logger.info(f"Generating session summary, message count: {len(messages)}")
            
            # Use the correct API call method
            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt
            )
            
            # Get the text from the response
            if hasattr(response, 'text'):
                summary = response.text
            elif hasattr(response, 'parts'):
                summary = ''.join([part.text for part in response.parts if hasattr(part, 'text')])
            else:
                summary = "Unable to generate summary, unknown API response format"
                
            logger.info(f"Summary generation successful, length: {len(summary)}")
            return summary
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            return "Unable to generate summary. The system may be experiencing temporary issues, please try again later."
    except Exception as e:
        logger.error(f"Error preparing summary prompt: {str(e)}")
        return "Unable to generate session summary."


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def delete_chat_session(request):
    """
    Completely delete the current active chat session, without saving any records
    Unlike end_chat_session, this function will delete the session and its messages from the database
    """
    try:
        # Get the current session ID
        session_id = request.session.get('current_chat_session_id')
        
        if not session_id:
            return Response(
                {"error": "No active chat session found"},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Find the session, ensuring only the user's own sessions can be deleted
        try:
            chat_session = ChatSession.objects.get(
                id=session_id,
                user=request.user
            )
        except ChatSession.DoesNotExist:
            return Response(
                {"error": "Chat session not found"},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Log the deletion operation
        logger.info(f"Deleting chat session {session_id} for user {request.user.username}")
        
        # First, delete all related messages
        ChatMessage.objects.filter(session=chat_session).delete()
        
        # Then, delete the session itself
        chat_session.delete()
        
        # Remove the session ID from the Django session
        request.session.pop('current_chat_session_id', None)
        
        return Response(
            {"success": True, "message": "Chat session completely deleted"},
            status=status.HTTP_200_OK
        )
    
    except Exception as e:
        logger.error(f"Error deleting chat session: {str(e)}")
        return Response(
            {"error": f"Error deleting chat session: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_chat_sessions(request):
    """Get the list of chat sessions for the user"""
    # Pagination parameters
    page = int(request.GET.get('page', 1))
    page_size = int(request.GET.get('page_size', 10))
    
    # Calculate the offset
    start = (page - 1) * page_size
    end = start + page_size
    
    # Get the user's session list
    sessions = ChatSession.objects.filter(user=request.user).order_by('-created_at')[start:end]
    
    # Count the total number of sessions
    total_count = ChatSession.objects.filter(user=request.user).count()
    
    # Prepare the response data
    session_list = []
    for session in sessions:
        session_data = {
            'id': session.id,
            'title': session.title,
            'created_at': session.created_at.strftime('%Y-%m-%d %H:%M'),
            'ended_at': session.ended_at.strftime('%Y-%m-%d %H:%M') if session.ended_at else None,
            'is_active': session.is_active,
            'duration': session.get_duration(),
            'videos_count': session.get_videos_count(),
            'messages_count': session.get_messages_count(),
            'has_summary': bool(session.summary)
        }
        session_list.append(session_data)
    
    return JsonResponse({
        'status': 'success',
        'sessions': session_list,
        'total': total_count,
        'page': page,
        'page_size': page_size,
        'pages': (total_count + page_size - 1) // page_size  # Calculate the total number of pages
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_chat_session(request, session_id):
    """Get the detailed information of a specific session"""
    try:
        session = ChatSession.objects.get(id=session_id, user=request.user)
        
        # Get the session messages
        messages = session.messages.all().order_by('timestamp')
        message_list = []
        for msg in messages:
            message_data = {
                'id': msg.id,
                'role': msg.role,
                'content': msg.content,
                'timestamp': msg.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
                'video_id': msg.video.id if msg.video else None
            }
            message_list.append(message_data)
        
        # Get the associated videos
        videos = session.session_videos.all()
        video_list = []
        for video in videos:
            video_data = {
                'id': video.id,
                'title': video.title,
                'url': video.url
            }
            video_list.append(video_data)
        
        # Prepare the response data
        session_data = {
            'id': session.id,
            'title': session.title,
            'created_at': session.created_at.strftime('%Y-%m-%d %H:%M'),
            'ended_at': session.ended_at.strftime('%Y-%m-%d %H:%M') if session.ended_at else None,
            'is_active': session.is_active,
            'summary': session.summary,
            'duration': session.get_duration(),
            'messages': message_list,
            'videos': video_list
        }
        
        return JsonResponse({
            'status': 'success',
            'session': session_data
        })
        
    except ChatSession.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': 'Session does not exist'
        }, status=404)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_session_videos(request):
    """
    Get the list of videos associated with the current active session
    """
    try:
        # Get the current user's active session
        active_session = ChatSession.objects.filter(
            user=request.user,
            is_active=True
        ).first()
        
        if not active_session:
            return Response(
                {"error": "No active session"},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Get the list of associated videos
        session_videos = Video.objects.filter(
            chat_session=active_session
        ).order_by('-created_at')  # Sort by creation time in reverse order
        
        # Format the video data
        videos_data = []
        for video in session_videos:
            videos_data.append({
                'id': video.id,
                'video_id': video.url.split('=')[-1] if '=' in video.url else video.url,  # Extract the YouTube video ID
                'title': video.title,
                'url': video.url,
                'added_at': video.created_at.strftime('%Y-%m-%d %H:%M'),
                'is_current': video.id == request.GET.get('current_video_id')  # Mark the current video
            })
        
        return Response({
            'session_id': active_session.id,
            'session_title': active_session.title,
            'created_at': active_session.created_at.strftime('%Y-%m-%d %H:%M'),
            'videos_count': len(videos_data),
            'videos': videos_data
        })
    
    except Exception as e:
        return Response(
            {"error": f"Failed to get session videos: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_chat_notes(request, session_id=None):
    """Export chat notes"""
    # If no session ID is provided, use the current active session
    if not session_id:
        session_id = request.session.get('current_chat_session_id')
        if not session_id:
            return JsonResponse({'status': 'error', 'message': 'No active session'})
    
    try:
        # Get the specified session
        session = ChatSession.objects.get(id=session_id, user=request.user)
        
        # Only return the session ID and user information, without generating an export file
        # Get basic information
        videos_count = session.session_videos.count()
        messages_count = session.messages.count()
        
        # Create a simplified response
        response = JsonResponse({
            'status': 'success',
            'message': 'Session ended',
            'session_id': session.id,
            'videos_count': videos_count,
            'messages_count': messages_count,
            'summary': session.summary or ''
        })
        
        logger.info(f"User {request.user.username} exported chat notes for session ID: {session.id}")
        return response
        
    except ChatSession.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Session does not exist'}, status=404)


# Web views for chat sessions
class ChatSessionListView(LoginRequiredMixin, ListView):
    """List view for chat sessions (notes)"""
    model = ChatSession
    template_name = 'api/notes.html'
    context_object_name = 'chat_sessions'
    paginate_by = 12
    
    def get_queryset(self):
        """Get chat sessions for the current user"""
        return ChatSession.objects.filter(user=self.request.user).order_by('-created_at')


class ChatSessionDetailView(LoginRequiredMixin, DetailView):
    """Detail view for a specific chat session (note)"""
    model = ChatSession
    template_name = 'api/notes_detail.html'
    context_object_name = 'chat_session'
    
    def get_object(self):
        """Get the chat session object"""
        session_id = self.kwargs.get('pk')
        return get_object_or_404(ChatSession, id=session_id, user=self.request.user)


from django.contrib.auth.decorators import login_required
@login_required
def delete_chat_session_web(request, session_id):
    """Delete a chat session from the web interface"""
    try:
        # Get the session
        session = get_object_or_404(ChatSession, id=session_id, user=request.user)
        
        # Delete the session and all associated messages
        session_id = session.id
        session.delete()
        
        # Add success message
        messages.success(request, f"聊天记录已成功删除")
        
        # Redirect to notes list
        return redirect('notes')
        
    except Exception as e:
        logger.error(f"Error deleting chat session: {str(e)}")
        messages.error(request, f"删除失败: {str(e)}")
        return redirect('notes_detail', pk=session_id)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_chat_session(request, session_id):
    """Download a chat session in the specified format"""
    # Get the format from the query parameter
    format_type = request.GET.get('format', 'txt')
    
    # Get the session
    session = get_object_or_404(ChatSession, id=session_id, user=request.user)
    
    # Get the messages
    messages_list = session.messages.all().order_by('timestamp')
    
    # Create a filename
    clean_title = ''.join(c if c.isalnum() or c in ' -_' else '_' for c in session.title)
    filename = f"DejaVocab_Chat_{clean_title}_{timezone.now().strftime('%Y%m%d')}"
    
    if format_type == 'txt':
        # Create a text file
        response = HttpResponse(content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
        
        # Write header
        response.write(f"标题: {session.title}\n")
        response.write(f"日期: {session.created_at.strftime('%Y-%m-%d %H:%M')}\n")
        response.write(f"消息数: {session.get_messages_count()}\n")
        response.write(f"持续时间: {session.get_duration()} 分钟\n\n")
        
        # Write summary if exists
        if session.summary:
            response.write("摘要:\n")
            response.write(f"{session.summary}\n\n")
        
        # Write messages
        response.write("聊天记录:\n")
        response.write("="*50 + "\n\n")
        
        for msg in messages_list:
            role_display = msg.get_role_display()
            response.write(f"{role_display} ({msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')}):\n")
            response.write(f"{msg.content}\n\n")
            response.write("-"*50 + "\n\n")
        
        return response
        
    elif format_type == 'docx':
        # Create a Word document
        doc = docx.Document()
        
        # Set title
        title = doc.add_heading(session.title, level=1)
        
        # Add metadata
        doc.add_paragraph(f"日期: {session.created_at.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"消息数: {session.get_messages_count()}")
        doc.add_paragraph(f"持续时间: {session.get_duration()} 分钟")
        
        # Add summary if exists
        if session.summary:
            doc.add_heading("摘要", level=2)
            doc.add_paragraph(session.summary)
        
        # Add messages
        doc.add_heading("聊天记录", level=2)
        
        for msg in messages_list:
            # Create a paragraph for each message
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"{msg.get_role_display()} ({msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')}):")
            role_run.bold = True
            
            # Add message content
            doc.add_paragraph(msg.content)
            
            # Add a separator
            doc.add_paragraph("---")
        
        # Save to a BytesIO object
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        
        # Return the document
        response = HttpResponse(f.read())
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response
        
    elif format_type == 'pdf':
        # Create a PDF document
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            fontSize=16,
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=10
        )
        
        normal_style = styles['Normal']
        
        role_style = ParagraphStyle(
            'Role',
            parent=normal_style,
            fontName='Helvetica-Bold',
            fontSize=10,
            spaceAfter=6
        )
        
        # Create content
        content = []
        
        # Add title
        content.append(Paragraph(session.title, title_style))
        content.append(Spacer(1, 10))
        
        # Add metadata
        content.append(Paragraph(f"日期: {session.created_at.strftime('%Y-%m-%d %H:%M')}", normal_style))
        content.append(Paragraph(f"消息数: {session.get_messages_count()}", normal_style))
        content.append(Paragraph(f"持续时间: {session.get_duration()} 分钟", normal_style))
        content.append(Spacer(1, 10))
        
        # Add summary if exists
        if session.summary:
            content.append(Paragraph("摘要", heading_style))
            content.append(Paragraph(session.summary, normal_style))
            content.append(Spacer(1, 10))
        
        # Add messages
        content.append(Paragraph("聊天记录", heading_style))
        content.append(Spacer(1, 10))
        
        for msg in messages_list:
            # Add role and timestamp
            content.append(Paragraph(f"{msg.get_role_display()} ({msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')}):", role_style))
            
            # Add message content
            content.append(Paragraph(msg.content, normal_style))
            content.append(Paragraph("---", normal_style))
            content.append(Spacer(1, 5))
        
        # Build the PDF
        doc.build(content)
        buffer.seek(0)
        
        # Return the PDF
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response
    
    else:
        # Default to text format if format is not recognized
        return HttpResponse("Unsupported format", status=400)
